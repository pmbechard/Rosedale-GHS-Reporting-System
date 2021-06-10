import mechanize
from bs4 import BeautifulSoup
import http.cookiejar
import re
from docx import Document
from datetime import date

# Authorization Config
url = 'https://ghs.rosedaleedu.com/login/index.php'
cj = http.cookiejar.CookieJar()
br = mechanize.Browser()
br.set_cookiejar(cj)
br.open(url)

courses = None
name = ''


def moodle_login(username, password):
    global name
    br.select_form(nr=0)
    br.form['username'] = username
    br.form['password'] = password
    br.submit()

    # Scraping for course list
    try:
        soup = BeautifulSoup(br.response().read(), 'lxml')
        name = soup.find('div', class_='myprofileitem fullname').text
        name = name.strip().split(' ')
        name = ' '.join(name[1:])
    except AttributeError:
        return False
    else:
        return f'Welcome, {name}!'


def course_choices():
    global courses
    soup = BeautifulSoup(br.response().read(), 'lxml')
    courses = soup.find_all('h3', class_='coursename')

    # Iteration of courses into dictionary
    counter = 1
    courses_dict = {}
    for course in courses:
        courses_dict[counter] = course.text.strip()
        counter += 1

    return courses_dict


def class_list_creation(course_choice):
    global courses
    # Searching dictionary for course
    for course in courses:
        if course_choice in course.text:
            # Find link for course page
            course_link = re.findall("http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+", str(course))
            course_link = course_link[0]

            # Opening course page
            br.open(course_link)
            soup = BeautifulSoup(br.response().read(), 'lxml')
            side_menu = soup.find_all('li', class_='r0')

            # Find link for active students page
            class_list_link = re.findall("http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\(\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+", str(side_menu[3]))
            class_list_link = class_list_link[0]

            # Opening active students page
            br.open(class_list_link)
            soup = BeautifulSoup(br.response().read(), 'lxml')
            class_list = soup.find_all('a', class_='', target='_blank')

            # Iterating list of students
            students = []
            for student in class_list:
                if not student.text.isdigit():
                    students.append(student.text.strip())

    return students


def comments_options(students, grades_dict):
    document = Document('CT Comments and Learning Skills.docx')
    available_comments = {}

    for student in students:
        available_comments[student] = []
        for i in range(6):
            grade = (grades_dict[students.index(student)][i]).upper().strip()
            if i == 0:
                for option in range(1, 13):
                    if grade == document.tables[2].cell(option, 1).text[0]:
                        available_comments[student].append(document.tables[2].cell(option, 1).text)
            if i == 1:
                for option in range(14, 22):
                    if grade == document.tables[2].cell(option, 1).text[0]:
                        available_comments[student].append(document.tables[2].cell(option, 1).text)
            if i == 2:
                for option in range(23, 31):
                    if grade == document.tables[2].cell(option, 1).text[0]:
                        available_comments[student].append(document.tables[2].cell(option, 1).text)
            if i == 3:
                for option in range(32, 40):
                    if grade == document.tables[2].cell(option, 1).text[0]:
                        available_comments[student].append(document.tables[2].cell(option, 1).text)
            if i == 4:
                for option in range(42, 50):
                    if grade == document.tables[2].cell(option, 1).text[0]:
                        available_comments[student].append(document.tables[2].cell(option, 1).text)
            if i == 5:
                for option in range(51, 67):
                    if grade == document.tables[2].cell(option, 1).text[0]:
                        available_comments[student].append(document.tables[2].cell(option, 1).text)

    return available_comments


def doc_creation(course_choice, students, term, grades_dict, comments):
    global name
    document = Document('CT Comments and Learning Skills.docx')

    course_choice = course_choice.split(' ')
    if course_choice[0][0] == '*':
        document.save(f'{course_choice[0][1:]} - CT Comments and Learning Skills - {date.today()}.docx')
    else:
        document.save(f'{course_choice[0]} - CT Comments and Learning Skills - {date.today()}.docx')

    # Fill in first table
    document.tables[0].cell(0, 1).text = 'Northeastern University branch - Shenyang'
    document.tables[0].cell(1, 1).text = course_choice[0][1:]
    document.tables[0].cell(2, 1).text = term
    document.tables[0].cell(3, 1).text = name
    document.tables[0].cell(4, 1).text = str(len(students))

    if course_choice[0][0] == '*':
        document.save(f'{course_choice[0][1:]} - CT Comments and Learning Skills - {date.today()}.docx')
    else:
        document.save(f'{course_choice[0]} - CT Comments and Learning Skills - {date.today()}.docx')

    # Check if second table's rows are enough, add more if not; insert student names into rows
    if len(document.tables[1].rows) < len(students):
        for i in range(len(students) - len(document.tables[1].rows) + 1):
            document.tables[1].add_row()

    # Add students names, grades and comments
    count = 1
    for student in students:
        document.tables[1].cell(count, 0).text = student
        for i in range(6):
            document.tables[1].cell(count, i+1).text = grades_dict[students.index(student)][i].upper().strip()
        for comment in comments[student]:
            for row in range(len(document.tables[2].rows)):
                if comment == document.tables[2].cell(row, 1).text:
                    if not document.tables[1].cell(count, 7).text:
                        document.tables[1].cell(count, 7).text = document.tables[2].cell(row, 0).text
                    else:
                        document.tables[1].cell(count, 8).text = document.tables[2].cell(row, 0).text
        count += 1

    if course_choice[0][0] == '*':
        document.save(f'{course_choice[0][1:]} - CT Comments and Learning Skills - {date.today()}.docx')
    else:
        document.save(f'{course_choice[0]} - CT Comments and Learning Skills - {date.today()}.docx')

